"""Ingest FinOps knowledge documents into the Chroma vector store.

Reads every ``.pdf`` / ``.docx`` / ``.md`` / ``.txt`` from ``KNOWLEDGE_DIR``,
extracts text, cleans and overlap-chunks it, embeds each chunk with the
configured provider, and upserts into the collection with metadata
``{source, tipo, chunk, char_len}``.

Run it manually (it is **not** started by the API or docker compose):

    python -m app.rag.ingest                 # ingest everything under KNOWLEDGE_DIR
    python -m app.rag.ingest --dry-run       # report what would be ingested, write nothing

The repository ships with an empty ``knowledge/`` directory — no documents are
loaded by default. Tenants drop their FinOps material in and run this command.
"""

from __future__ import annotations

import argparse
import asyncio
import logging
import sys
from pathlib import Path

from app.config import settings
from app.rag import embeddings
from app.rag.chunking import chunk_text, clean_text, detect_tipo
from app.rag.store import reset_collection

logger = logging.getLogger(__name__)
BATCH = 64  # chunks embedded per round-trip


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in (".md", ".txt"):
        return _read_text(path)
    raise ValueError(f"unsupported file type: {suffix}")


def _iter_docs(root: Path) -> list[Path]:
    return sorted(
        p for p in root.rglob("*")
        if p.is_file() and p.suffix.lower() in (".pdf", ".docx", ".md", ".txt")
    )


async def _ingest(dry_run: bool) -> dict:
    root = Path(settings.knowledge_dir)
    if not root.exists():
        raise FileNotFoundError(f"knowledge dir not found: {root}")
    docs = _iter_docs(root)
    if not docs:
        logger.warning("no documents found in %s — nothing to ingest", root)
        return {"documents": 0, "chunks": 0, "dry_run": dry_run}

    collection = None if dry_run else reset_collection()
    total_chunks = 0
    next_id = 0

    for doc in docs:
        try:
            text = clean_text(_read(doc))
        except Exception as exc:
            logger.warning("skip %s: %s", doc.name, exc)
            continue
        if not text:
            logger.warning("skip %s: no extractable text", doc.name)
            continue
        chunks = chunk_text(text)
        tipo = detect_tipo(doc.name)
        logger.info("→ %s [%s] %d chunks", doc.name, tipo, len(chunks))
        total_chunks += len(chunks)
        if dry_run:
            continue

        for i in range(0, len(chunks), BATCH):
            batch = chunks[i:i + BATCH]
            vectors = await embeddings.embed_texts(batch)
            ids = [f"doc{next_id + j}" for j in range(len(batch))]
            metadatas = [
                {
                    "source": doc.name,
                    "tipo": tipo,
                    "chunk": i + j,
                    "char_len": len(batch[j]),
                }
                for j in range(len(batch))
            ]
            collection.add(ids=ids, embeddings=vectors, documents=batch, metadatas=metadatas)
            next_id += len(batch)

    return {"documents": len(docs), "chunks": total_chunks, "dry_run": dry_run}


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    parser = argparse.ArgumentParser(description="Ingest FinOps knowledge into the RAG store")
    parser.add_argument("--dry-run", action="store_true", help="report only, write nothing")
    args = parser.parse_args()
    try:
        result = asyncio.run(_ingest(args.dry_run))
    except FileNotFoundError as exc:
        print(f"❌ {exc}", file=sys.stderr)
        sys.exit(1)
    print(f"✅ ingested {result['documents']} documents / {result['chunks']} chunks"
          f"{' (dry-run)' if result['dry_run'] else ''}")


if __name__ == "__main__":
    main()
